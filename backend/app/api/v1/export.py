"""
8D Report Word export - Professional format.
"""
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from datetime import datetime
import tempfile
import json
import os

from app.database.session import get_db
from app.models.user import User
from app.models.quality_case import QualityCase
from app.core.security import require_user

router = APIRouter()

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn("w:shd"), {
        qn("w:fill"): color,
        qn("w:val"): "clear",
    })
    shading_elm.append(shading)


def add_section(doc, d_number, title, content):
    """Add a formatted D-section to the document."""
    # Section heading
    p = doc.add_paragraph()
    run = p.add_run(f"{d_number} - {title}")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(5, 28, 44)  # mckinsey-navy
    p.space_after = Pt(6)

    # Content
    if content and content.strip():
        for line in content.strip().split("\n"):
            cp = doc.add_paragraph(line.strip())
            cp.paragraph_format.space_after = Pt(4)
            cp.paragraph_format.left_indent = Cm(0.5)
            for run in cp.runs:
                run.font.size = Pt(10)
    else:
        cp = doc.add_paragraph("(待完善)")
        cp.paragraph_format.left_indent = Cm(0.5)
        cp.runs[0].font.size = Pt(10)
        cp.runs[0].font.color.rgb = RGBColor(139, 157, 175)  # muted

    # Separator line
    doc.add_paragraph().space_after = Pt(2)


@router.get("/cases/{case_id}/export-8d")
async def export_8d_word(
    case_id: int,
    user: User = Depends(require_user),
    db: Session = Depends(get_db),
):
    """Export 8D report as professionally formatted Word document."""
    if not HAS_DOCX:
        raise HTTPException(status_code=500, detail="python-docx not installed")

    case = db.query(QualityCase).filter(QualityCase.id == case_id, QualityCase.user_id == user.id).first()
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(10)

    # === Header ===
    header_para = doc.add_paragraph()
    header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = header_para.add_run("8D PROBLEM SOLVING REPORT")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(5, 28, 44)
    header_para.space_after = Pt(4)

    # Subtitle
    sub = doc.add_paragraph()
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = sub.add_run("AI Quality Portal - Auto Generated")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(139, 157, 175)
    sub.space_after = Pt(20)

    # === Info table ===
    info_table = doc.add_table(rows=3, cols=4)
    info_table.style = "Table Grid"
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Row 1
    info_table.cell(0, 0).text = "Report No."
    info_table.cell(0, 1).text = f"8D-{case.id:04d}"
    info_table.cell(0, 2).text = "Date"
    info_table.cell(0, 3).text = datetime.now().strftime("%Y-%m-%d")

    # Row 2
    info_table.cell(1, 0).text = "Case Type"
    type_labels = {"complaint": "Customer Complaint", "incoming": "Incoming Material",
                   "process": "Process Abnormality", "failure": "Product Failure",
                   "supplier": "Supplier Issue", "internal": "Internal Quality Issue"}
    info_table.cell(1, 1).text = type_labels.get(case.case_type, case.case_type)
    info_table.cell(1, 2).text = "Owner"
    info_table.cell(1, 3).text = user.nickname or user.email or "Quality Engineer"

    # Row 3
    info_table.cell(2, 0).text = "Title"
    info_table.cell(2, 1).text = case.title or "Untitled"
    info_table.cell(2, 2).text = "Status"
    info_table.cell(2, 3).text = case.status.upper() if case.status else "OPEN"

    # Format info table
    for row in info_table.rows:
        for i, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(4)
                paragraph.paragraph_format.space_after = Pt(4)
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            if i % 2 == 0:  # Label cells
                set_cell_shading(cell, "F5F7FA")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

    doc.add_paragraph().space_after = Pt(16)

    # === D Sections ===
    # Parse measures
    containment = ""
    corrective = ""
    preventive = ""
    if case.measures:
        try:
            m = json.loads(case.measures)
            containment = m.get("containment", "")
            corrective = m.get("corrective", "")
            preventive = m.get("preventive", "")
        except (json.JSONDecodeError, TypeError):
            corrective = case.measures

    add_section(doc, "D0", "Emergency Response Action (ERA)",
                containment or "Refer to D3 containment measures.")

    add_section(doc, "D1", "Team Formation",
                f"Team Lead: {user.nickname or user.email}\nGenerated by: AI Quality Portal")

    add_section(doc, "D2", "Problem Description",
                case.problem_statement)

    add_section(doc, "D3", "Interim Containment Action (ICA)",
                containment)

    add_section(doc, "D4", "Root Cause Analysis",
                case.root_cause)

    add_section(doc, "D5", "Permanent Corrective Action (PCA)",
                corrective)

    add_section(doc, "D6", "Verification of Effectiveness",
                "Pending verification - to be completed after corrective action implementation.")

    add_section(doc, "D7", "Prevention of Recurrence",
                preventive)

    add_section(doc, "D8", "Team Recognition & Closure",
                "Pending final review and closure.")

    # === Footer ===
    doc.add_paragraph().space_after = Pt(20)
    footer = doc.add_paragraph()
    footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = footer.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')} | AI Quality Portal | aidmaic.top")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(139, 157, 175)

    # Save
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)
    tmp.close()

    filename = f"8D_Report_{case.id:04d}_{datetime.now().strftime('%Y%m%d')}.docx"
    return FileResponse(
        tmp.name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename,
    )
